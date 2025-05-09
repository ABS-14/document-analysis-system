import io
import os
import re
import mimetypes

# Define supported file types
supported_file_types = {
    "txt": "text/plain",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
    "rtf": "application/rtf",
}

def extract_text_from_file(uploaded_file):
    """
    Extract text from various file formats.
    
    Args:
        uploaded_file: A file object from Streamlit file_uploader
        
    Returns:
        str: The extracted text content
        
    Raises:
        ValueError: If the file format is not supported
    """
    # Get the file extension
    file_extension = os.path.splitext(uploaded_file.name)[1].lower().replace(".", "")
    
    # Check if file type is supported
    if file_extension not in supported_file_types:
        raise ValueError(f"Unsupported file format: {file_extension}. Supported formats are: {', '.join(supported_file_types.keys())}")
    
    # Process based on file type
    if file_extension == "txt":
        # For text files, just read the content
        return uploaded_file.getvalue().decode("utf-8")
    
    elif file_extension == "pdf":
        # For PDF files, use PyPDF2 with error handling
        try:
            import PyPDF2
            
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
            
            # Extract text from each page
            text = ""
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n\n"
            
            return text
        except ImportError:
            raise ValueError("PDF extraction requires PyPDF2 package. Please install it or use text input instead.")
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    elif file_extension in ["docx", "doc"]:
        # For DOCX files, use python-docx with error handling
        try:
            import docx
            
            # Load the document
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            
            # Extract text from paragraphs
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"
            
            return text
        except ImportError:
            raise ValueError("DOCX extraction requires python-docx package. Please install it or use text input instead.")
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    elif file_extension == "rtf":
        # For RTF files, use striprtf with error handling
        try:
            from striprtf.striprtf import rtf_to_text
            
            # Read and convert RTF content
            rtf_text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
            plain_text = rtf_to_text(rtf_text)
            
            return plain_text
        except ImportError:
            raise ValueError("RTF extraction requires striprtf package. Please install it or use text input instead.")
        except Exception as e:
            raise ValueError(f"Error extracting text from RTF: {str(e)}")
    
    else:
        # This should not happen due to earlier check
        raise ValueError(f"Unsupported file format: {file_extension}")

def clean_text(text):
    """
    Clean and normalize text for processing.
    
    Args:
        text (str): The input text
        
    Returns:
        str: Cleaned text
    """
    # Replace multiple spaces with a single space
    text = re.sub(r'\s+', ' ', text)
    
    # Remove extra newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Trim whitespace
    text = text.strip()
    
    return text
